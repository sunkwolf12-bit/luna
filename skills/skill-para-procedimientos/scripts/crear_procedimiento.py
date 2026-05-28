#!/usr/bin/env python3
"""
Crea un procedimiento institucional a partir del archivo base y un Word fuente.
Uso: python3 crear_procedimiento.py <archivo_fuente.docx> <CLAVE> <NOMBRE> <FECHA> [DEPARTAMENTO]

Ejemplo:
  python3 crear_procedimiento.py PCD1_seguimiento.docx PCD1 "PROCEDIMIENTO DE SEGUIMIENTO A COBRANZA ATRASADA" "25/10/2025" "Cobranza"

Salida: workspaces/instructivos/CLAVE_NOMBRE_SANITIZADO.docx

IMAGENES Y TABLAS: se conservan automaticamente del documento fuente al DESARROLLO.
"""
import sys, os, re, shutil, zipfile, unicodedata, io
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from lxml import etree
from copy import deepcopy as dc

WORKSPACE = '/home/elena/.openclaw/workspace'
BASE_FILE = os.path.join(WORKSPACE, 'workspaces/instructivos/FORMATO_INSTITUCIONAL_PROCEDIMIENTO_2025.docx')
OUT_DIR = os.path.join(WORKSPACE, 'workspaces/instructivos')

# Namespaces usados
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
PKG_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'


def slug(text):
    """Convierte texto a un nombre de archivo seguro (sin acentos, mayusculas, solo A-Z0-9_)"""
    s = ''.join(c for c in unicodedata.normalize('NFD', text) if unicodedata.category(c) != 'Mn')
    return re.sub(r'[^A-Z0-9]+', '_', s.upper()).strip('_')


def insert_paragraph_after(paragraph, text):
    """Inserta un nuevo parrafo con 'text' justo despues de 'paragraph'."""
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def patch_header(docx_path, nombre, fecha, clave, departamento):
    """Modifica header2.xml y document.xml directamente en el ZIP."""
    with zipfile.ZipFile(docx_path, 'r') as z:
        files = {n: z.read(n) for n in z.namelist()}

    # --- DOCUMENT.XML: portada ---
    files['word/document.xml'] = files['word/document.xml'].replace(
        b'PROCEDIMIENTO EN BLANCO', nombre.encode('utf-8'))

    # --- HEADER2.XML ---
    h2 = files.get('word/header2.xml') or files.get('word/header1.xml', b'')
    if not h2:
        print("ERROR: No se encontro header2.xml ni header1.xml en el archivo base")
        return

    # 1) Titulo en encabezado
    h2 = h2.replace(b'PROCEDIMIENTO EN BLANCO', nombre.encode('utf-8'))

    # 2) Fecha
    h2 = h2.replace(b'31/08/202</w:t></w:r><w:r w:rsidR="00BF75E2"><w:t>5</w:t>',
                    (fecha + '</w:t>').encode())

    # 3) CLAVE - insertar DESPUES de "CLAVE: " en la MISMA celda
    h2 = re.sub(
        rb'(<w:t[^>]*>CLAVE:\s*)(</w:t>)',
        ('\\1' + clave + '\\2').encode(),
        h2
    )

    # 4) DEPARTAMENTO
    dep_pos = h2.find(b'DEPARTAMENTO:')
    ppr = h2.find(b'<w:pPr>', dep_pos + 100)
    ppr_end = h2.find(b'</w:pPr>', ppr)
    para_end = h2.find(b'</w:p>', ppr_end)
    if ppr > 0 and para_end > ppr:
        between = h2[ppr_end + len(b'</w:pPr>'):para_end]
        if between.strip() == b'':
            insert2 = ('<w:r><w:rPr><w:b/><w:bCs/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:t>'
                       + departamento + '</w:t></w:r>').encode()
            h2 = h2[:para_end] + insert2 + h2[para_end:]

    files['word/header2.xml'] = h2

    # --- ELIMINAR placeholders, dataBindings y arreglar core.xml ---
    for key in list(files.keys()):
        if key.endswith('.xml'):
            files[key] = re.sub(rb'<w:placeholder>.*?</w:placeholder>', b'', files[key])
            files[key] = re.sub(rb'<w:dataBinding[^>]*/>', b'', files[key])
            files[key] = re.sub(rb'<w:dataBinding>.*?</w:dataBinding>', b'', files[key], flags=re.DOTALL)
            files[key] = files[key].replace(b'<w:showingPlcHdr/>', b'')
            files[key] = files[key].replace(b'<w:showingPlcHdr />', b'')

    # Fix docProps/core.xml
    if 'docProps/core.xml' in files:
        files['docProps/core.xml'] = files['docProps/core.xml'].replace(
            b'<dc:subject>PROCEDIMIENTO EN BLANCO</dc:subject>',
            f'<dc:subject>{nombre}</dc:subject>'.encode('utf-8'))

    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)


def fill_body_initial(docx_path, objetivo, alcance, docref_lines, def_lines):
    """Rellena OBJETIVO, ALCANCE, DOC REF y DEF. El DESARROLLO se llena despues via inyeccion directa."""
    d = Document(docx_path)
    labels = ['OBJETIVO:', 'ALCANCE:', 'DOCUMENTOS DE REFERENCIA:', 'DEFINICIONES:', 'DESARROLLO']
    for label_name in reversed(labels[:4]):  # Solo los primeros 4
        cur_i = None
        for j, pp in enumerate(d.paragraphs):
            if pp.text.strip().upper() == label_name:
                cur_i = j
                break
        if cur_i is None:
            continue
        nxt = None
        for j in range(cur_i + 1, len(d.paragraphs)):
            if d.paragraphs[j].text.strip().upper() in labels:
                nxt = j
                break
        if nxt is None:
            nxt = len(d.paragraphs)
        for _ in range(nxt - cur_i - 1):
            rm = d.paragraphs[cur_i + 1]
            rm._element.getparent().remove(rm._element)

    def find(label):
        for p in d.paragraphs:
            if p.text.strip().upper() == label:
                return p
        raise ValueError(f'Label no encontrado: {label}')

    cur = find('OBJETIVO:')
    insert_paragraph_after(cur, objetivo)
    cur = find('ALCANCE:')
    insert_paragraph_after(cur, alcance)
    cur = find('DOCUMENTOS DE REFERENCIA:')
    last = cur
    for line in docref_lines:
        last = insert_paragraph_after(last, line)
    cur = find('DEFINICIONES:')
    last = cur
    for line in def_lines:
        last = insert_paragraph_after(last, line)

    d.save(docx_path)


def inject_desarrollo_from_source(output_path, fuente_path, start_marker='3.'):
    """
    Reemplaza el contenido DESPUES del label DESARROLLO en output_path con
    el body del documento fuente (desde start_marker como '3. Responsables').
    Preserva: parrafos, tablas, imagenes, SDTs (structured document tags).
    Copia archivos media (imagenes) y remapea sus rIds.
    """
    with zipfile.ZipFile(output_path, 'r') as z:
        out_files = {n: z.read(n) for n in z.namelist()}
    with zipfile.ZipFile(fuente_path, 'r') as z:
        src_files = {n: z.read(n) for n in z.namelist()}

    # Parsear XMLs
    src_xml = etree.fromstring(src_files['word/document.xml'])
    out_xml = etree.fromstring(out_files['word/document.xml'])
    src_body = src_xml.find(f'{{{W}}}body')
    out_body = out_xml.find(f'{{{W}}}body')
    


    # Encontrar label DESARROLLO en output
    desarrollo = None
    for p in out_body.findall(f'{{{W}}}p'):
        texts = []
        for t in p.iter(f'{{{W}}}t'):
            if t.text:
                texts.append(t.text)
        full_text = ''.join(texts).strip().upper()
        if full_text == 'DESARROLLO':
            desarrollo = p
            break

    if desarrollo is None:
        print("WARNING: label DESARROLLO no encontrado en el output")
        return

    # Encontrar contenido a copiar desde fuente (desde start_marker)
    encontrado = False
    src_elements = []
    for child in src_body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        # Detectar inicio (ej. "3. Responsables")
        if tag == 'p':
            texts = []
            for t in child.iter(f'{{{W}}}t'):
                if t.text:
                    texts.append(t.text)
            full_text = ''.join(texts)
            if start_marker in full_text and not encontrado:
                encontrado = True
                src_elements.append(child)
                continue

        if encontrado and tag in ('p', 'tbl', 'sdt'):
            src_elements.append(child)
        elif encontrado and tag == 'sectPr':
            continue  # sectPr es metadata, no contenido

    if not src_elements:
        print(f"WARNING: no se encontro contenido desde '{start_marker}' en fuente")
        return

    # --- Reconstruir body completo: mantener todo hasta DESARROLLO, luego inyectar fuente, luego sectPr ---
    # En vez de insertar en medio (que tiene bugs con lxml), construimos una lista nueva y reemplazamos.
    out_body_children = list(out_body)
    dev_idx = None
    for i, child in enumerate(out_body_children):
        if child is desarrollo:
            dev_idx = i
            break
    
    # Extraer sectPr (guardarlo antes de limpiar)
    sectPr_elem = None
    for sp in out_body.findall(f'{{{W}}}sectPr'):
        sectPr_elem = sp
        out_body.remove(sp)
    
    # Limpiar todo
    for child in list(out_body):
        out_body.remove(child)
    
    # Reconstruir: elementos antes de DESARROLLO + DESARROLLO + source + sectPr
    for child in out_body_children[:dev_idx]:
        out_body.append(child)
    out_body.append(desarrollo)
    for elem in src_elements:
        out_body.append(dc(elem))
    if sectPr_elem is not None:
        out_body.append(sectPr_elem)
    

    
    # --- Gestionar imagenes: remapear rIds ---
    src_rels = {}
    src_rels_path = 'word/_rels/document.xml.rels'
    if src_rels_path in src_files:
        rels_xml = etree.fromstring(src_files[src_rels_path])
        for rel in rels_xml:
            tag = rel.tag.split('}')[-1] if '}' in rel.tag else rel.tag
            if tag == 'Relationship':
                r_id = rel.get('Id')
                target = rel.get('Target', '')
                rtype = rel.get('Type', '')
                if 'image' in rtype.lower() or 'media' in target.lower():
                    src_rels[r_id] = (target, rtype)

    # Determinar siguiente rId disponible en output
    out_rels_path = 'word/_rels/document.xml.rels'
    next_num = 100  # empezar alto para no colisionar con rIds de la plantilla base
    if out_rels_path in out_files:
        out_rels_xml = etree.fromstring(out_files[out_rels_path])
        for rel in out_rels_xml:
            tag = rel.tag.split('}')[-1] if '}' in rel.tag else rel.tag
            if tag == 'Relationship':
                rid = rel.get('Id', '')
                nums = re.findall(r'\d+', rid)
                if nums:
                    next_num = max(next_num, int(nums[0]) + 1)

    # Mapear rIds viejos -> nuevos
    rid_map = {}
    for old_rid, (target, rtype) in src_rels.items():
        new_rid = f'rId{next_num}'
        next_num += 1
        rid_map[old_rid] = (new_rid, target, rtype)

    # Remapear rIds en los elementos insertados
    if rid_map:
        for elem in src_elements:
            for el in elem.iter():
                for attr_name in [
                    f'{{{REL_NS}}}embed',
                    f'{{{REL_NS}}}id',
                    f'{{{REL_NS}}}link',
                ]:
                    val = el.get(attr_name)
                    if val and val in rid_map:
                        el.set(attr_name, rid_map[val][0])

    # --- Copiar archivos media ---
    for old_rid, (new_rid, target, rtype) in rid_map.items():
        basename = target.split('/')[-1]  # ej. "image1.png"
        # Encontrar el archivo en source
        src_media_name = None
        candidates = [f'word/media/{basename}', f'word/{target}']
        if target.startswith('media/'):
            candidates.insert(0, f'word/{target}')

        for candidate in candidates:
            if candidate in src_files:
                src_media_name = candidate
                break

        if src_media_name is None:
            # Busqueda por nombre de archivo parcial
            for name in src_files:
                if name.endswith(basename):
                    src_media_name = name
                    break

        if src_media_name:
            out_media_name = f'word/media/{basename}'
            # Si ya existe en output, usar un nombre unico
            if out_media_name in out_files:
                base, ext = os.path.splitext(basename)
                out_media_name = f'word/media/{base}_src{next_num}{ext}'
                # Actualizar el rId en los elementos
                for elem in src_elements:
                    for el in elem.iter():
                        for attr_name in [f'{{{REL_NS}}}embed', f'{{{REL_NS}}}id']:
                            if el.get(attr_name) == new_rid:
                                pass  # keep same rId, media file just renamed
            out_files[out_media_name] = src_files[src_media_name]

    # --- Agregar relaciones de imagenes a output ---
    if rid_map and out_rels_path in out_files:
        out_rels_xml = etree.fromstring(out_files[out_rels_path])
        for old_rid, (new_rid, target, rtype) in rid_map.items():
            new_rel = etree.SubElement(out_rels_xml, f'{{{PKG_NS}}}Relationship')
            new_rel.set('Id', new_rid)
            new_rel.set('Type', rtype)
            new_rel.set('Target', f'media/{target.split("/")[-1]}')
        out_files[out_rels_path] = etree.tostring(out_rels_xml, xml_declaration=True, encoding='UTF-8')

    # --- Actualizar Content_Types si hay imagenes nuevas ---
    ctypes_path = '[Content_Types].xml'
    if rid_map and ctypes_path in out_files:
        ctypes_xml = etree.fromstring(out_files[ctypes_path])
        existing = set()
        for ov in ctypes_xml:
            tag = ov.tag.split('}')[-1] if '}' in ov.tag else ov.tag
            if tag == 'Override':
                existing.add(ov.get('PartName', ''))

        for old_rid, (new_rid, target, rtype) in rid_map.items():
            basename = target.split('/')[-1]
            part_name = f'/word/media/{basename}'
            if part_name not in existing:
                import mimetypes
                mime, _ = mimetypes.guess_type(basename)
                if mime is None:
                    ext = os.path.splitext(basename)[1].lower()
                    mime_map = {
                        '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                        '.gif': 'image/gif', '.bmp': 'image/bmp', '.tiff': 'image/tiff',
                    }
                    mime = mime_map.get(ext, 'image/png')
                existing.add(part_name)
                new_ov = etree.SubElement(ctypes_xml, f'{{{ctypes_xml.nsmap[None]}}}Override')
                new_ov.set('PartName', part_name)
                new_ov.set('ContentType', mime)

        out_files[ctypes_path] = etree.tostring(ctypes_xml, xml_declaration=True, encoding='UTF-8')

    # --- Guardar document.xml ---
    out_files['word/document.xml'] = etree.tostring(out_xml, xml_declaration=True, encoding='UTF-8')

    # --- Escribir ZIP final ---
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in out_files.items():
            zout.writestr(name, data)

    print(f"  Inyectados {len(src_elements)} elementos al DESARROLLO "
          f"(imagenes copiadas: {len(rid_map)})")


def main():
    if len(sys.argv) < 5:
        print("Uso: python3 crear_procedimiento.py <fuente.docx> <CLAVE> <NOMBRE> <FECHA> [DEPARTAMENTO]")
        sys.exit(1)

    fuente = sys.argv[1]
    clave = sys.argv[2].strip()
    nombre = sys.argv[3].strip()
    fecha = sys.argv[4].strip()
    departamento = sys.argv[5].strip() if len(sys.argv) > 5 else 'Cobranza'

    if not os.path.exists(fuente):
        print(f"ERROR: Archivo fuente no encontrado: {fuente}")
        sys.exit(1)
    if not os.path.exists(BASE_FILE):
        print(f"ERROR: Archivo base no encontrado: {BASE_FILE}")
        sys.exit(1)

    out_name = f"{clave}_{slug(nombre)}.docx"
    out_path = os.path.join(OUT_DIR, out_name)

    # --- PASO 1: Extraer contenido del archivo fuente (TEXTO para OBJ/ALC) ---
    src = Document(fuente)
    sp = [p.text.rstrip() for p in src.paragraphs if p.text.strip()]

    def idx(pattern):
        for i, t in enumerate(sp):
            if re.match(pattern, t.strip(), re.I):
                return i
        return None

    i_obj = idx(r'^1\.?\s*OBJETIVO')
    i_alc = idx(r'^2\.?\s*ALCANCE')

    objetivo = sp[i_obj + 1] if i_obj is not None and i_obj + 1 < len(sp) else ''
    alcance = sp[i_alc + 1] if i_alc is not None and i_alc + 1 < len(sp) else ''

    # Generar DOCUMENTOS DE REFERENCIA
    docref_generados = [
        'CRM PROTEG-RT y/o SIGA (segun aplique al procedimiento).',
        'Archivos internos de respaldo mencionados en el procedimiento.'
    ]
    for line in sp:
        up = line.upper()
        if 'CRM' in up:
            docref_generados[0] = 'CRM PROTEG-RT (mencionado en el procedimiento).'
        if 'SIGA' in up:
            docref_generados[0] = 'SIGA y CRM PROTEG-RT (mencionados en el procedimiento).'
        if 'PAGOS_V' in up or 'PAGOS V' in up:
            docref_generados[1] = 'Archivo Pagos_V3 (segun referencia en el procedimiento).'

    defs_generados = ['Definiciones basadas en el contenido del procedimiento.']

    # --- PASO 2: Copiar base, parchar headers, llenar secciones iniciales ---
    shutil.copyfile(BASE_FILE, out_path)
    patch_header(out_path, nombre, fecha, clave, departamento)
    fill_body_initial(out_path, objetivo, alcance, docref_generados, defs_generados)

    # --- PASO 3: Inyectar DESARROLLO desde fuente (preserva tablas, imagenes) ---
    inject_desarrollo_from_source(out_path, fuente, start_marker='3.')

    print(f"Procedimiento creado: {out_path}")
    print(f"  CLAVE: {clave}")
    print(f"  NOMBRE: {nombre}")
    print(f"  FECHA: {fecha}")
    print(f"  DEPARTAMENTO: {departamento}")


if __name__ == '__main__':
    main()
